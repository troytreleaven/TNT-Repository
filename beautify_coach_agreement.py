#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Read the markdown content
with open('/data/.openclaw/workspace/Independent_Coach_Agreement_Draft.md', 'r') as f:
    content = f.read()

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Define colors
dark_blue = RGBColor(0, 51, 102)
maroon = RGBColor(128, 0, 0)
dark_gray = RGBColor(64, 64, 64)

# Helper function to set cell shading
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# TITLE
title = doc.add_heading('INDEPENDENT COACH AGREEMENT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = dark_blue
    run.font.bold = True
    run.font.size = Pt(20)

# Subtitle
subtitle = doc.add_paragraph('Dale Carnegie Business Group & Independent Coach')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = dark_gray
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.size = Pt(12)

subtitle2 = doc.add_paragraph('Ontario Massage Therapist Association Engagement')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.color.rgb = dark_gray
subtitle2.runs[0].font.size = Pt(11)

doc.add_paragraph()

# Add horizontal line
doc.add_paragraph('_' * 80)
doc.add_paragraph()

# Process content line by line
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # Skip title lines (already added)
    if line.startswith('# INDEPENDENT COACH AGREEMENT') or line.startswith('## Dale Carnegie'):
        i += 1
        continue
    
    # Section headers (##)
    if line.startswith('## '):
        heading_text = line[3:].strip()
        heading = doc.add_heading(heading_text, level=1)
        for run in heading.runs:
            run.font.color.rgb = dark_blue
            run.font.bold = True
        doc.add_paragraph()
    
    # Subsection headers (###)
    elif line.startswith('### '):
        heading_text = line[4:].strip()
        heading = doc.add_heading(heading_text, level=2)
        for run in heading.runs:
            run.font.color.rgb = maroon
            run.font.bold = True
    
    # Bold text (**text**)
    elif line.startswith('**') and line.endswith('**') and len(line) > 4:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line[2:-2])
        run.bold = True
        run.font.color.rgb = dark_gray
    
    # List items (- or [ ])
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        # Check for bold within list item
        if '**' in text:
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold parts
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(text, style='List Bullet')
    
    # Checkbox items
    elif line.startswith('- [ ]'):
        text = line[5:].strip()
        p = doc.add_paragraph(f'☐ {text}', style='List Bullet')
    
    # Regular paragraph
    elif line and not line.startswith('---') and not line.startswith('|'):
        # Skip empty or formatting lines
        if not line.startswith('**RECOMMENDATION**'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
    
    # Tables (simplified - just add as text)
    elif line.startswith('|') and '---' not in line:
        # Skip table formatting lines
        pass
    
    # Horizontal rules
    elif line.startswith('---'):
        doc.add_paragraph('_' * 80)
    
    # Page breaks before signatures
    elif line.startswith('## SIGNATURES'):
        doc.add_page_break()
        heading = doc.add_heading('SIGNATURES', level=1)
        for run in heading.runs:
            run.font.color.rgb = dark_blue
    
    # Signature blocks
    elif 'By:' in line or 'Name:' in line or 'Title:' in line or 'Date:' in line:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(12)
    
    i += 1

# Add signature lines at the end
doc.add_paragraph()
doc.add_paragraph('_' * 80)
doc.add_paragraph()

# Save document
output_path = '/data/.openclaw/media/outbound/Independent_Coach_Agreement.docx'
doc.save(output_path)
print(f"✅ Document created: {output_path}")
