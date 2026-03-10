#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Read the markdown content
with open('/data/.openclaw/workspace/Coach_Partnership_Executive_Summary.md', 'r') as f:
    content = f.read()

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Define colors
dark_blue = RGBColor(0, 51, 102)
maroon = RGBColor(128, 0, 0)
dark_gray = RGBColor(64, 64, 64)

# TITLE
title = doc.add_heading('COACH PARTNERSHIP AGREEMENT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = dark_blue
    run.font.bold = True
    run.font.size = Pt(20)

# Subtitle
subtitle = doc.add_paragraph('Executive Summary for Initial Discussions')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = dark_gray
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('_' * 80)
doc.add_paragraph()

# Process content line by line
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # Skip title lines (already added)
    if line.startswith('# COACH PARTNERSHIP AGREEMENT') or line.startswith('## Executive Summary'):
        i += 1
        continue
    
    # Section headers (##)
    if line.startswith('## '):
        heading_text = line[3:].strip()
        heading = doc.add_heading(heading_text, level=1)
        for run in heading.runs:
            run.font.color.rgb = dark_blue
            run.font.bold = True
        doc.add_paragraph()
    
    # Subsection headers (###)
    elif line.startswith('### '):
        heading_text = line[4:].strip()
        heading = doc.add_heading(heading_text, level=2)
        for run in heading.runs:
            run.font.color.rgb = maroon
            run.font.bold = True
    
    # Table separator
    elif line.startswith('|---'):
        i += 1
        continue
    
    # Table rows
    elif line.startswith('| ') and '---' not in line:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if len(cells) == 2:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run1 = p.add_run(f"{cells[0]}: ")
            run1.bold = True
            p.add_run(cells[1])
    
    # List items
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        doc.add_paragraph(text, style='List Bullet')
    
    # Bold text (**text**)
    elif line.startswith('**') and line.endswith('**') and len(line) > 4:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line[2:-2])
        run.bold = True
        run.font.color.rgb = dark_gray
    
    # Regular paragraph
    elif line and not line.startswith('---') and not line.startswith('|'):
        if 'A:' in line or '**' in line:
            p = doc.add_paragraph(line)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
    
    # Horizontal rules
    elif line.startswith('---'):
        doc.add_paragraph('_' * 80)
    
    i += 1

# Save document
output_path = '/data/.openclaw/media/outbound/Coach_Partnership_Executive_Summary.docx'
doc.save(output_path)
print(f"✅ Executive Summary created: {output_path}")
