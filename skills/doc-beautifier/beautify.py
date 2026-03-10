#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Define colors
dark_blue = RGBColor(0, 51, 102)
gold = RGBColor(218, 165, 32)
dark_gray = RGBColor(64, 64, 64)

# Helper function to set shading
def set_shading(run, color):
    r = run._element
    r.rPr = r.get_or_add_rPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    r.rPr.append(shading_elm)

# TITLE
title = doc.add_heading('The AI Tsunami: Is Your Team Ready to Swim?', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = dark_blue
    run.font.bold = True

# Subtitle
subtitle = doc.add_paragraph('Canadians Leading with Impact Podcast')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = dark_gray
subtitle.runs[0].font.bold = True

subtitle2 = doc.add_paragraph('Episode __ | 60 Minutes')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.color.rgb = dark_gray

doc.add_paragraph()

# SECTION 1: OPENING
h1 = doc.add_heading('OPENING', 1)
h1.runs[0].font.color.rgb = dark_blue

doc.add_paragraph('⏱️ 5 minutes', style='Intense Quote')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Quick relatable scenario: ').bold = True
p.add_run('"Imagine telling your team in 6 months that AI just did in 4 hours what used to take your company 3 years and $50 million."')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Introduce the thesis: ').bold = True
p.add_run('The AI tsunami is coming — but it makes human leadership MORE valuable, not less.')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Set the Canadian context: ').bold = True
p.add_run('"This isn\'t just a US story. It\'s happening in Canadian boardrooms right now."')

doc.add_page_break()

# SECTION 2: THE TSUNAMI
h1 = doc.add_heading('SECTION 1: THE TSUNAMI IS HERE', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 12 minutes', style='Intense Quote')
doc.add_paragraph()

# Global Stats Table
doc.add_heading('Global Stats', 2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Stat'
hdr_cells[1].text = 'Source'
hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True

data = [
    ('300 million jobs globally could be affected by AI', 'Goldman Sachs'),
    ('57% of U.S. work hours can be automated', 'McKinsey'),
    ('6-7% of U.S. workforce displaced if AI widely adopted', 'Goldman Sachs'),
    ('88% of organizations now use AI', 'McKinsey'),
]

for i, (stat, source) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = stat
    row[1].text = source

doc.add_paragraph()

# Canadian Stats Table
doc.add_heading('Canadian Stats', 2)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Stat'
hdr_cells2[1].text = 'Source'
hdr_cells2[0].paragraphs[0].runs[0].bold = True
hdr_cells2[1].paragraphs[0].runs[0].bold = True

data2 = [
    ('60% of Canadian workers in high AI-exposure occupations', 'Vector Institute'),
    ('12% of Canadian businesses now use AI (doubled from 6% in 2024)', 'Statistics Canada'),
    ('29% of Canadian workers use AI weekly', 'Business Insider'),
    ('2.4% of Canadian job postings mention GenAI', 'SI Systems'),
]

for i, (stat, source) in enumerate(data2):
    row = table2.rows[i+1].cells
    row[0].text = stat
    row[1].text = source

doc.add_paragraph()

# Quote
quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run('"Intelligence is the new electricity. The grid is live, the voltage is climbing, and the meter is running."').italic = True
quote.add_run(' — Solve Everything').bold = True

doc.add_page_break()

# SECTION 3: WHAT MACHINES CAN'T DO
h1 = doc.add_heading('SECTION 2: WHAT THE MACHINES CAN\'T DO', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 12 minutes', style='Intense Quote')
doc.add_paragraph()

# Expert Quotes
doc.add_heading('Expert Quotes', 2)

quotes = [
    ('Forbes', 'AI is swallowing routine work. What remains—and what now differentiates leaders—are enduring human skills like emotional intelligence.'),
    ('World Economic Forum', 'The abilities to ask better questions, navigate ambiguity, empathize and turn ideas into action are not teachable by machines.'),
    ('Deidre Lipton, Microsoft Canada', 'AI should make work more human.'),
]

for source, quote_text in quotes:
    p = doc.add_paragraph()
    p.add_run(f'"{quote_text}"').italic = True
    p.add_run(f' — {source}').bold = True
    doc.add_paragraph()

# Four things list
doc.add_heading('The Four Things Robots Can\'t Do', 2)

things = [
    'Build trust with a struggling team member',
    'Navigate a difficult conversation with empathy',
    'Inspire people to follow through uncertainty',
    'Create a culture where humans actually want to work',
]

for i, thing in enumerate(things, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(thing)

doc.add_page_break()

# SECTION 4: SHIFT
h1 = doc.add_heading('SECTION 3: FROM EMPLOYEE TO ENTREPRENEUR & CREATOR', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 10 minutes', style='Intense Quote')
doc.add_paragraph()

quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run('"The future isn\'t the employee. It\'s the entrepreneur and the creator. These are the two things that will stand apart from AI."').italic = True

doc.add_heading('Why This Matters', 2)

reasons = [
    'AI can execute, analyze, optimize',
    'AI can\'t create something from nothing',
    'AI can\'t take risks, bet on themselves, or build something from zero',
    'That\'s the human advantage',
]

for reason in reasons:
    p = doc.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run(reason)

doc.add_page_break()

# SECTION 5: OPPORTUNITY
h1 = doc.add_heading('SECTION 4: THE OPPORTUNITY FOR CANADIAN LEADERS', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 10 minutes', style='Intense Quote')
doc.add_paragraph()

points = [
    ('Canadian companies are behind', 'Only 12% using AI (vs. 88% globally). This is OUR advantage if we move first.'),
    ('Leadership skills are the differentiator', 'WEF: "Agency now matters more than information"'),
    ('Soft skills = hard business results', 'TalentSmartEQ 2026: human skills are the strongest predictor of organizational performance'),
    ('Skill gaps', 'Are the #1 barrier to AI adoption (WTW 2026)'),
]

for title, desc in points:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# Callout box
callout = doc.add_paragraph()
callout.style = 'Intense Quote'
callout.add_run('The Conversation Every Leader Needs to Have:').bold = True
doc.add_paragraph()
callout2 = doc.add_paragraph()
callout2.add_run('"If you\'re not using AI and constantly learning how to use AI tools, we won\'t have a role for you. I don\'t want you doing mundane admin tasks. I want you connecting with humans, engaging personally, and leaving the routine to the machines."')

doc.add_page_break()

# SECTION 6: ACTION ITEMS
h1 = doc.add_heading('SECTION 5: WHAT SHOULD CANADIAN LEADERS DO?', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 8 minutes', style='Intense Quote')
doc.add_paragraph()

doc.add_heading('Three Actions', 2)

actions = [
    ('Invest in EQ', 'Emotional intelligence is your moat'),
    ('Model the way', 'Show vulnerability, lead through change'),
    ('Build culture', 'Trust, connection, purpose (what AI can\'t replicate)'),
]

for action, desc in actions:
    p = doc.add_paragraph()
    p.add_run(f'{action} — ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# Dale Carnegie Angle
h1 = doc.add_heading('The Dale Carnegie Angle', 2)
quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run('"You\'re not selling training. You\'re buying their future."').italic = True
quote.add_run(' — Dale Carnegie').bold = True

doc.add_page_break()

# CLOSING
h1 = doc.add_heading('CLOSING', 1)
h1.runs[0].font.color.rgb = dark_blue
doc.add_paragraph('⏱️ 3 minutes', style='Intense Quote')
doc.add_paragraph()

quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run('"The leaders who thrive won\'t be the ones who out-work the machines. They\'ll be the ones who out-connect with their people."').italic = True

doc.add_paragraph()

quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run('"The future isn\'t the employee. It\'s the entrepreneur and the creator."').italic = True

doc.add_paragraph()

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Episode __ | Canadians Leading with Impact | Dale Carnegie Training — Greater Toronto Area')
footer_run.font.color.rgb = dark_gray
footer_run.font.size = Pt(10)

# Save
doc.save('/data/.openclaw/workspace/podcast-ai-tsunami-beautiful.docx')
print('Document created successfully!')
