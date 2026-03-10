#!/usr/bin/env python3
"""
Doc Beautifier - Markdown to Word Converter
Converts markdown files to professionally formatted Word documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

def set_shading(run, color):
    """Set background shading for a run"""
    r = run._element
    r.rPr = r.get_or_add_rPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    r.rPr.append(shading_elm)

def parse_markdown_table(lines, start_idx):
    """Parse a markdown table and return data + end index"""
    headers = []
    rows = []
    
    # Header row
    header_line = lines[start_idx].strip()
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    
    # Skip separator line
    idx = start_idx + 2
    
    # Data rows
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_line = lines[idx].strip()
        cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
        idx += 1
    
    return headers, rows, idx

def convert_markdown_to_docx(input_file, output_file):
    """Convert markdown file to Word document"""
    
    # Read markdown content
    with open(input_file, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Define colors
    dark_blue = RGBColor(0, 51, 102)
    gold = RGBColor(218, 165, 32)
    dark_gray = RGBColor(64, 64, 64)
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Page break marker
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue
        
        # H1 headings
        if line.startswith('# '):
            text = line[2:].strip()
            h1 = doc.add_heading(text, 1)
            h1.runs[0].font.color.rgb = dark_blue
            i += 1
            continue
        
        # H2 headings
        if line.startswith('## '):
            text = line[3:].strip()
            h2 = doc.add_heading(text, 2)
            h2.runs[0].font.color.rgb = dark_blue
            i += 1
            continue
        
        # H3 headings
        if line.startswith('### '):
            text = line[4:].strip()
            h3 = doc.add_heading(text, 3)
            h3.runs[0].font.color.rgb = dark_gray
            i += 1
            continue
        
        # Tables
        if line.strip().startswith('|') and '|' in line[1:]:
            headers, rows, end_idx = parse_markdown_table(lines, i)
            
            # Create table
            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            for j, header in enumerate(headers):
                hdr_cells[j].text = header
                for run in hdr_cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = dark_blue
            
            # Data rows
            for row_idx, row_data in enumerate(rows):
                cells = table.rows[row_idx+1].cells
                for j, cell_text in enumerate(row_data):
                    if j < len(cells):
                        # Clean up markdown formatting
                        cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)  # Bold
                        cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)      # Italic
                        cells[j].text = cell_text
            
            doc.add_paragraph()
            i = end_idx
            continue
        
        # Blockquotes
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            # Collect multi-line quotes
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_text += ' ' + lines[i].strip()[1:].strip()
                i += 1
            
            quote = doc.add_paragraph()
            quote.style = 'Quote'
            quote.add_run(quote_text).italic = True
            continue
        
        # Bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Clean up markdown formatting in bullet
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            i += 1            
            # Handle nested bullets
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                text = lines[i].strip()[2:]
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(text)
                i += 1
            continue
        
        # Numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if numbered_match:
            text = numbered_match.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            i += 1
            continue
        
        # Regular paragraphs with formatting
        text = line.strip()
        
        # Skip table separator lines
        if re.match(r'^\|[-:\s|]+\|$', text):
            i += 1
            continue
        
        # Clean up markdown formatting
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italic
        text = re.sub(r'`(.+?)`', r'\1', text)        # Code
        
        # Handle inline formatting for actual bold/italic
        p = doc.add_paragraph()
        
        # Simple parsing for bold text
        parts = re.split(r'(\*\*.+?\*\*)', line.strip())
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f'Document saved to: {output_file}')

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        if not output_file.endswith('.docx'):
            output_file += '.docx'
    elif len(sys.argv) == 2:
        input_file = sys.argv[1]
        output_file = input_file.rsplit('.', 1)[0] + '.docx'
    else:
        input_file = '/data/.openclaw/workspace/research/cold-email-system-research.md'
        output_file = '/data/.openclaw/media/outbound/Cold_Email_System_Research_Report.docx'
    
    convert_markdown_to_docx(input_file, output_file)
